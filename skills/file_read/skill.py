from __future__ import annotations

import re
from pathlib import Path
from typing import Any


class FileReadSkill:
    """Read local files (DOC/DOCX/TXT) and extract text content."""

    @staticmethod
    def read(path: str | Path) -> dict[str, Any]:
        p = Path(path)
        if not p.exists():
            return {"error": f"文件不存在: {p}"}

        suffix = p.suffix.lower()
        handlers: dict[str, Any] = {
            ".pdf": FileReadSkill._read_pdf,
            ".docx": FileReadSkill._read_docx,
            ".doc": FileReadSkill._read_doc,
            ".txt": FileReadSkill._read_txt,
        }

        handler = handlers.get(suffix)
        if handler is None:
            return {"error": f"不支持的文件类型: {suffix}"}

        try:
            text, metadata = handler(p)
            return {
                "doc_type": suffix.lstrip("."),
                "path": str(p),
                "text": text,
                "metadata": metadata,
            }
        except Exception as e:
            return {"error": f"读取失败: {e}", "path": str(p)}

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    @staticmethod
    def _read_pdf(path: Path) -> tuple[str, dict]:
        import fitz

        doc = fitz.open(str(path))

        parts: list[str] = []
        meta = doc.metadata or {}
        total_tables = 0

        for i, page in enumerate(doc, start=1):
            parts.append(f"--- 第 {i} 页 ---")
            parts.append(page.get_text().strip())

            # 尝试提取表格
            try:
                tables = page.find_tables()
                for table in tables:
                    total_tables += 1
                    parts.append("")
                    for row in table.extract():
                        parts.append(" | ".join(str(c).strip() if c else "" for c in row))
            except Exception:
                pass

        text = "\n".join(parts).strip()
        pages = doc.page_count

        # 收集元数据
        metadata = {
            "pages": pages,
            "characters": len(text),
        }
        for key in ("title", "author", "subject", "keywords"):
            val = meta.get(key, "")
            if val:
                metadata[key] = val

        doc.close()
        return text, metadata

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    @staticmethod
    def _read_docx(path: Path) -> tuple[str, dict]:
        from docx import Document

        doc = Document(str(path))

        parts: list[str] = []
        for p in doc.paragraphs:
            parts.append(p.text)

        for table in doc.tables:
            parts.append("")
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))

        text = "\n".join(parts).strip()
        return text, {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "characters": len(text),
        }

    # ------------------------------------------------------------------
    # DOC (legacy binary format)
    # ------------------------------------------------------------------

    @staticmethod
    def _read_doc(path: Path) -> tuple[str, dict]:
        # 有些 .doc 实质上是 OOXML 格式，先试 python-docx
        try:
            return FileReadSkill._read_docx(path)
        except Exception:
            pass

        # Windows 下优先用 Win32 COM（需安装 Microsoft Word）
        try:
            return FileReadSkill._read_doc_win32com(path)
        except Exception:
            pass

        # 纯 Python fallback：olefile 解析二进制流
        return FileReadSkill._read_doc_olefile(path)

    @staticmethod
    def _read_doc_win32com(path: Path) -> tuple[str, dict]:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False

            doc = word.Documents.Open(str(path.absolute()), ReadOnly=True)
            text = doc.Content.Text

            # 清理 Word 插入的控制字符
            text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
            text = re.sub(r"\r\n?", "\n", text).strip()
            return text, {"source": "win32com", "characters": len(text)}
        finally:
            if doc is not None:
                try:
                    doc.Close(SaveChanges=False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    @staticmethod
    def _read_doc_olefile(path: Path) -> tuple[str, dict]:
        import struct

        import olefile

        with olefile.OleFileIO(str(path)) as ole:
            if not ole.exists("WordDocument"):
                raise ValueError("不是有效的 Word 文档")

            data = ole.openstream("WordDocument").read()
            # FIB.ccpText — 文档主体字符数（UTF-16 code units）
            ccp_text = struct.unpack_from("<I", data, 0x004C)[0]

            # 文本通常从偏移 2048 附近开始
            best_off = 2048
            raw = data[best_off : best_off + ccp_text * 2].decode(
                "utf-16-le", errors="replace"
            )
            cleaned = "".join(c for c in raw if c.isprintable() or c in "\n\r\t")
            cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", cleaned)
            cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
            return cleaned, {"source": "olefile", "characters": len(cleaned)}

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    @staticmethod
    def _read_txt(path: Path) -> tuple[str, dict]:
        for encoding in ["utf-8", "gbk", "gb2312", "utf-16"]:
            try:
                text = path.read_text(encoding=encoding)
                return text.strip(), {"encoding": encoding, "characters": len(text)}
            except (UnicodeDecodeError, UnicodeError):
                continue
        text = path.read_text(encoding="utf-8", errors="ignore")
        return text.strip(), {"encoding": "utf-8 (with errors ignored)", "characters": len(text)}
