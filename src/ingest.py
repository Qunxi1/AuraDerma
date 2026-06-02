from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

from llm import LLMClient
from schema import DocType, DocumentChunk, ProductRecord, new_id


@dataclass(slots=True)
class IngestedDocument:
    doc_id: str
    path: Path
    doc_type: DocType
    chunks: list[DocumentChunk]


def detect_doc_type(path: Path) -> DocType:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return DocType.PDF
    if suffix in {".docx", ".doc"}:
        return DocType.DOCX
    return DocType.TXT


def read_document(path: Path) -> str:
    doc_type = detect_doc_type(path)
    if doc_type == DocType.PDF:
        with fitz.open(path) as pdf:
            return "\n".join(page.get_text("text") for page in pdf)
    if doc_type == DocType.DOCX:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    return path.read_text(encoding="utf-8", errors="ignore")


def chunk_text(text: str, chunk_size: int = 900, overlap: int = 120) -> list[str]:
    text = " ".join(text.split())
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = max(0, end - overlap)
    return chunks


def ingest_document(path: str | Path) -> IngestedDocument:
    path = Path(path)
    text = read_document(path)
    doc_type = detect_doc_type(path)
    doc_id = new_id()
    chunks = [
        DocumentChunk(
            chunk_id=new_id(),
            doc_id=doc_id,
            doc_type=doc_type,
            text=chunk,
            title=path.stem,
            section="",
            page=None,
            chunk_index=i,
            source_uri=str(path),
        )
        for i, chunk in enumerate(chunk_text(text))
    ]
    return IngestedDocument(doc_id=doc_id, path=path, doc_type=doc_type, chunks=chunks)


def ingest_product_rows(rows: list[dict], llm: LLMClient) -> list[ProductRecord]:
    records: list[ProductRecord] = []
    for row in rows:
        records.append(
            ProductRecord(
                product_id=row.get("product_id") or new_id(),
                name=row["name"],
                brand=row.get("brand", ""),
                category=row.get("category", ""),
                price_cny=row.get("price_cny"),
                price_note=row.get("price_note"),
                ingredients=row.get("ingredients", []),
                ingredient_ordered_text=row.get("ingredient_ordered_text", ""),
                skin_types=row.get("skin_types", []),
                concerns=row.get("concerns", []),
                usage_notes=row.get("usage_notes", ""),
                source=row.get("source", ""),
            )
        )
    return records
